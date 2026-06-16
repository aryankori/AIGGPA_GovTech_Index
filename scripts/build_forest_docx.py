"""
Build a premium, publication-ready DOCX assessment report for the Forest Department
based on the 40 surveyed employees in AIGGPA_Forest_Department_Report.xlsx.
"""
import os, subprocess, json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from openpyxl import load_workbook

GOG = r"C:\Users\aryan\.gemini\antigravity\bin\gog.exe"
ACCOUNT = "aryan.kori14@gmail.com"

# ── Color Palette (Forest Green & Gold Theme) ──
FOREST_DARK = RGBColor(0x1B, 0x5E, 0x20)  # Primary Header Color
FOREST_MED  = RGBColor(0x2E, 0x7D, 0x32)  # Secondary Accent Color
GOLD_ACCENT = RGBColor(0xC4, 0x9A, 0x2A)  # Decorative Line Color
CHARCOAL    = RGBColor(0x21, 0x21, 0x21)  # Body Text
MUTED_GRAY  = RGBColor(0x66, 0x66, 0x66)  # Metadata & Captions

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def build_docx_report():
    print("Loading survey workbook...")
    wb = load_workbook("AIGGPA_Forest_Department_Report.xlsx")
    ws_list = wb["Respondent List"]
    ws_analysis = wb["Class-wise Analysis"]
    
    # Read respondents
    respondents = []
    for r in range(2, 42):
        row_vals = [ws_list.cell(row=r, column=c).value for c in range(1, 13)]
        respondents.append(row_vals)
    print(f"Loaded {len(respondents)} respondents from Excel list.")
    
    # Read Class-wise Analysis
    analysis_data = []
    for r in range(1, 40):
        m_name = ws_analysis.cell(row=r, column=1).value
        c1 = ws_analysis.cell(row=r, column=2).value
        c2 = ws_analysis.cell(row=r, column=3).value
        c3 = ws_analysis.cell(row=r, column=4).value
        c4 = ws_analysis.cell(row=r, column=5).value
        if m_name:
            analysis_data.append((m_name, c1, c2, c3, c4))
            
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = CHARCOAL

    # ── Style Helpers ──
    def add_para(text, bold=False, italic=False, size=11, color=CHARCOAL, spacing_after=6, spacing_before=0, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        p.paragraph_format.space_after = Pt(spacing_after)
        p.paragraph_format.space_before = Pt(spacing_before)
        p.paragraph_format.line_spacing = 1.15
        if align: 
            p.alignment = align
        return p

    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        for run in h.runs:
            run.font.color.rgb = FOREST_DARK if level == 1 else FOREST_MED
            run.font.name = 'Calibri'
            run.bold = True
        return h

    def add_bullet_styled(text, bold_prefix=None, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.4 * level)
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.bold = True
            r1.font.color.rgb = CHARCOAL
            r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = CHARCOAL
        return p

    # ═════════════════════════════════════════════════════════════
    #  COVER PAGE (Forest Green Premium Layout)
    # ═════════════════════════════════════════════════════════════
    for _ in range(3): 
        doc.add_paragraph()
        
    add_para("AIGGPA GovTech Assessment Index", bold=True, size=13, color=MUTED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", bold=True, size=11, color=GOLD_ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_para("FOREST DEPARTMENT SURVEY REPORT", bold=True, size=26, color=FOREST_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
    
    add_para("Digital Adoption, Infrastructure, and Capacity Building Assessment\n(वन विभाग, मध्य प्रदेश - कर्मचारी डिजिटल आकलन)", 
             italic=True, size=12, color=MUTED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=18)
             
    add_para("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", bold=True, size=11, color=GOLD_ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    for _ in range(6): 
        doc.add_paragraph()
        
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    
    details = [
        ("Department:", "Madhya Pradesh Forest Department (वन विभाग)"),
        ("Primary Sample:", "40 Employees (10 respondents per Class 1, 2, 3, and 4)"),
        ("Methodology:", "In-Person Interviews & UTAUT Assessment Framework"),
        ("Prepared By:", "Aryan Kori, Intern"),
        ("Institution:", "Atal Bihari Vajpayee Institute of Good Governance and Policy Analysis (AIGGPA)")
    ]
    
    for idx, (label, val) in enumerate(details):
        r_cells = info_table.rows[idx].cells
        r_cells[0].text = label
        r_cells[0].paragraphs[0].runs[0].bold = True
        r_cells[0].paragraphs[0].runs[0].font.color.rgb = FOREST_DARK
        r_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_cells[1].text = val
        r_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for cell in r_cells:
            cell.width = Inches(3.2)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("Table of Contents", level=1)
    doc.add_paragraph().add_run("━" * 58).font.color.rgb = GOLD_ACCENT
    
    toc_items = [
        ("1. Executive Summary", "3"),
        ("2. The UTAUT Evaluation Framework", "4"),
        ("3. Demographics and Survey Structure", "4"),
        ("4. Class-wise Comparative Analysis", "5"),
        ("5. Key Research Findings", "6"),
        ("6. Detailed Respondent Profile (40 surveyed employees)", "7"),
        ("7. Cadre-Specific Actionable Recommendations", "9"),
    ]
    
    for title, page_no in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(title)
        r1.font.size = Pt(11)
        r1.font.color.rgb = CHARCOAL
        
        # Add dot leaders
        dots_count = 80 - len(title) - len(page_no)
        r2 = p.add_run("." * dots_count)
        r2.font.color.rgb = MUTED_GRAY
        
        r3 = p.add_run(page_no)
        r3.font.size = Pt(11)
        r3.font.color.rgb = CHARCOAL
        
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("1. Executive Summary", level=1)
    
    summary_text = (
        "This report outlines primary research findings regarding the digital readiness, usage, and capacity gaps "
        "among employees in the Madhya Pradesh Forest Department (वन विभाग). Modernizing forest administration "
        "requires frontline and administrative personnel to move from manual record-keeping to unified digital tools. "
        "This assessment covers 40 structured, in-person employee interviews, equally distributed (10 per group) "
        "across administrative classes: Class 1 (Senior IFS Officers), Class 2 (Assistant Directors/Range Officers), "
        "Class 3 (Accountants/Assistants), and Class 4 (Peons/Support Staff)."
    )
    add_para(summary_text)
    
    summary_text_2 = (
        "Key takeaways highlight a clear digital divide between administrative cadres. Senior officers (Class 1 and 2) "
        "show near-unanimous enthusiasm for digital platforms, maintaining 100% agreement on their value. "
        "However, frontline support staff (Class 4) experience a digital gap: 80% have received zero formal "
        "digital training, and their digital footprint is confined almost entirely to personal communications via WhatsApp. "
        "This gap impacts efficiency, highlighting the need for role-based training programs and conversational "
        "interfaces to build digital capacity across all departmental layers."
    )
    add_para(summary_text_2)

    # ═════════════════════════════════════════════════════════════
    #  2. THE UTAUT ASSESSMENT FRAMEWORK
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("2. The UTAUT Evaluation Framework", level=1)
    
    utaut_intro = (
        "To evaluate technology adoption systematically, this study uses the Unified Theory of Acceptance and Use "
        "of Technology (UTAUT) model. This framework assesses digital tools through four core constructs:"
    )
    add_para(utaut_intro)
    
    add_bullet_styled("The degree to which an employee believes that using a digital tool will help them perform job tasks faster and with higher accuracy. (e.g., e-Office reducing file movement delay).", bold_prefix="Performance Expectancy (PE): ")
    add_bullet_styled("The perceived difficulty or ease of learning and using departmental applications. Older administrative cadres and Class 4 staff often report high effort expectancy.", bold_prefix="Effort Expectancy (EE): ")
    add_bullet_styled("The influence of peer networks, formal departmental mandates, and senior encouragements in driving daily adoption.", bold_prefix="Social Influence (SI): ")
    add_bullet_styled("The availability of physical infrastructure, stable internet connectivity, on-premise IT helpdesks, and structured training programs.", bold_prefix="Facilitating Conditions (FC): ")

    # ═════════════════════════════════════════════════════════════
    #  3. DEMOGRAPHICS AND SURVEY STRUCTURE
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("3. Demographics and Survey Structure", level=1)
    
    demo_text = (
        "To ensure balanced insights, the data represents an equal split across four employee cadres. "
        "Each class represents 10 respondents, providing a reliable comparison of operational challenges "
        "at different levels of the organization."
    )
    add_para(demo_text)
    
    # Draw simple table of sample structure
    headers = ["Class Group", "Cadre Type", "Sample Count", "Primary Roles Coverd", "Typical Tech Usage Profile"]
    rows = [
        ["Class 1", "IFS / Senior Officers", "10", "DCF, CF, CCF, Deputy Directors", "High-end devices, e-Office, specialized GIS"],
        ["Class 2", "ACFs / Range Officers", "10", "Assistant Directors, Rangers", "Desktop, Mobile, Field tracking & GPS"],
        ["Class 3", "Clerical & Finance", "10", "Accountants, Assistants, Dispatch", "Desktop heavy, Web portals, PFMS, drafting"],
        ["Class 4", "Support Staff", "10", "Peons, Drivers, Daftaris", "Smartphone only, WhatsApp, file transit"]
    ]
    
    table = doc.add_table(rows=1+len(rows), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header styling
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "1B5E20")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for r_idx, row_data in enumerate(rows):
        add_cant_split(table.rows[r_idx+1])
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8F5E9")
                
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════
    #  4. CLASS-WISE COMPARATIVE ANALYSIS
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("4. Class-wise Comparative Analysis", level=1)
    
    analysis_intro = (
        "The survey reveals stark operational variations when data is cross-tabulated across "
        "the four administrative classes. Key averages and primary responses are outlined in the comparison table below:"
    )
    add_para(analysis_intro)
    
    # Filter and add important rows from Class-wise Analysis
    comp_headers = ["Metric Evaluated", "Class 1 (Senior)", "Class 2 (Field)", "Class 3 (Office)", "Class 4 (Support)"]
    comp_rows = []
    for item in analysis_data:
        # Filter down to the most important metrics to keep the document highly readable
        name = item[0]
        if any(keyword in name for keyword in ["Count", "Q10", "Q11", "Q12", "Q15", "Q16", "Q22", "Q27", "Q44", "Q75"]):
            comp_rows.append(item)
            
    comp_table = doc.add_table(rows=1+len(comp_rows), cols=5)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_table.style = 'Table Grid'
    
    hdr_cells = comp_table.rows[0].cells
    for i, h in enumerate(comp_headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "2E7D32")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for r_idx, row_data in enumerate(comp_rows):
        add_cant_split(comp_table.rows[r_idx+1])
        for c_idx, val in enumerate(row_data):
            cell = comp_table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if "──" in str(val):
                run.bold = True
                run.font.color.rgb = FOREST_DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F1F8E9")
                
    add_para("")

    # ═════════════════════════════════════════════════════════════
    #  5. KEY RESEARCH FINDINGS
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("5. Key Research Findings", level=1)
    
    add_bullet_styled(
        "While 100% of Class 1-3 employees agree that digital tools should be adopted to improve transparency, "
        "there is a significant drop in Class 4 where technical understanding is low.",
        bold_prefix="Positive Adoption Intent: "
    )
    add_bullet_styled(
        "Frontline peons (Class 4) report a score of 5/5 on difficulty (very difficult) and 1/5 on confidence (not confident). "
        "Many have never operated a desktop computer and rely solely on mobile phones for communication.",
        bold_prefix="The Support Cadre Barrier: "
    )
    add_bullet_styled(
        "Over 80% of Class 4 staff have received zero formal training. In contrast, Class 1-2 officers receive regular "
        "executive updates but note that standard training does not accommodate their schedules.",
        bold_prefix="Uneven Training Coverage: "
    )
    add_bullet_styled(
        "WhatsApp and basic drafting tools are widely used across all levels. However, specialized administrative tools "
        "(e.g., e-Office, GIS, IFMS) are limited to Class 1-3 personnel.",
        bold_prefix="Fragmented Tool Awareness: "
    )
    add_bullet_styled(
        "Headquarters offices have high internet bandwidth and stability (4-5/5), while field offices (Class 2 and 3) "
        "frequently report slow connectivity and power outages.",
        bold_prefix="Infrastructure Gaps in Field Offices: "
    )
    
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════
    #  6. DETAILED RESPONDENT PROFILE
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("6. Detailed Respondent Profile (40 Surveyed Employees)", level=1)
    
    grid_intro = (
        "The complete roster of the 40 surveyed Forest Department employees, including designation, "
        "cadre, demographics, and key digital indicator responses (Q15 and Q27), is detailed below:"
    )
    add_para(grid_intro)
    
    grid_headers = ["Name", "Designation", "Class", "Age", "Edu", "Difficulty (Q15)", "% Digital (Q27)"]
    grid_table = doc.add_table(rows=1+len(respondents), cols=7)
    grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid_table.style = 'Table Grid'
    
    hdr_cells = grid_table.rows[0].cells
    for i, h in enumerate(grid_headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "1B5E20")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for r_idx, row_data in enumerate(respondents):
        add_cant_split(grid_table.rows[r_idx+1])
        # Columns to print: Name (index 2), Designation (3), Class (4), Age (6), Edu (9), Q15 (10), Q27 (11)
        # Note: row_data has indices based on Excel list
        # Let's map accurately:
        # row_data[2]: Name
        # row_data[3]: Designation
        # row_data[4]: Class
        # row_data[6]: Age Group
        # row_data[9]: Education
        # row_data[10]: Q15
        # row_data[11]: Q27
        vals_to_print = [
            row_data[2],  # Name
            row_data[3],  # Designation
            row_data[4],  # Class
            row_data[6],  # Age
            row_data[9],  # Education
            row_data[10], # Q15
            row_data[11]  # Q27
        ]
        
        for c_idx, val in enumerate(vals_to_print):
            cell = grid_table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val if val is not None else "-")
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            if c_idx in [2, 5, 6]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8F5E9")
                
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════
    #  7. CADRE-SPECIFIC ACTIONABLE RECOMMENDATIONS
    # ═════════════════════════════════════════════════════════════
    add_heading_styled("7. Cadre-Specific Actionable Recommendations", level=1)
    
    rec_intro = (
        "To bridge the digital gap effectively, the department should transition from a "
        "one-size-fits-all training model to cadre-specific enablement programs:"
    )
    add_para(rec_intro)
    
    # Subsections for recommendations
    add_heading_styled("Class 1 & Class 2 (Senior Leadership & Field Managers)", level=2)
    add_bullet_styled(
        "Implement condensed, 45-minute executive webinars focusing on decision dashboards, "
        "compliance portals, and spatial tools (e-Green Watch, GIS), structured to fit busy schedules.",
        bold_prefix="Flexible Professional Enablement: "
    )
    add_bullet_styled(
        "Equip field-based Range Officers with department-issued rugged GPS units and offline GIS apps, "
        "reducing reliance on personal mobile data in remote forest areas.",
        bold_prefix="Infrastructure Support for Field Staff: "
    )
    
    add_heading_styled("Class 3 (Administrative Clerks & Finance Staff)", level=2)
    add_bullet_styled(
        "Design regular, hands-on workshops on IFMS, e-Office file tracking, and drafting shortcuts "
        "to increase speed and accuracy in daily tasks.",
        bold_prefix="Hands-on Digital Workshops: "
    )
    add_bullet_styled(
        "Simplify the user interfaces of internal portals and provide step-by-step checklists "
        "to minimize entry errors.",
        bold_prefix="Simplify Portal Interfaces: "
    )
    add_bullet_styled(
        "Establish an instant-response IT support chat or portal to quickly resolve connectivity or login issues.",
        bold_prefix="Dedicated Administrative IT Support: "
    )
    
    add_heading_styled("Class 4 (Support & Field Assistants)", level=2)
    add_bullet_styled(
        "Conduct foundational digital literacy workshops covering smartphone operations, e-Office tracking, "
        "and basic cyber hygiene.",
        bold_prefix="Foundational Digital Literacy: "
    )
    add_bullet_styled(
        "Integrate basic work functions with familiar tools (e.g., creating a secure WhatsApp bot for Class 4 "
        "staff to check file transit status without navigating complex portals).",
        bold_prefix="Conversational Tech Interfaces: "
    )
    add_bullet_styled(
        "Acknowledge the primary non-digital support role of Class 4 staff, ensuring they are not "
        "disadvantaged by administrative digital transitions.",
        bold_prefix="Maintain Essential Non-Digital Workflows: "
    )
    
    # Save document
    out_path = "AIGGPA_Forest_Department_Report.docx"
    doc.save(out_path)
    print(f"\n✅ DOCX saved: {out_path}")
    
    # Upload to Google Drive
    print("\nUploading DOCX to Google Drive...")
    cmd = [GOG, "--account", ACCOUNT, "--no-input", "--json",
           "drive", "upload", out_path]
    r = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
    if r.returncode == 0:
        data = json.loads(r.stdout)
        link = data.get("file", {}).get("webViewLink", "")
        print(f"   ✅ DOCX Uploaded successfully! Link: {link}")
        return link
    else:
        print(f"   Upload error: {r.stderr[:200]}")
        return None

if __name__ == "__main__":
    build_docx_report()
