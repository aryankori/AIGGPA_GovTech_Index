"""
Fill the CHR-T Interns Timeline DOCX for Aryan Kori
based on the existing Himani Singh template.
Replaces name, title, overview, and all daily entries.
"""
from docx import Document
import copy, re

src = r"c:\Users\aryan\Downloads\CHR-T_Interns_Timeline_April_2026_Final.docx"
dst = r"c:\Users\aryan\Downloads\CHR-T_Interns_Timeline_Aryan_Kori.docx"

doc = Document(src)

# ── Header paragraphs ──────────────────────────────────────────
for p in doc.paragraphs:
    if "Himani Singh" in p.text:
        # Won't appear in paragraphs, it's in the table
        pass

# ── Table 0: Weeks 1-2  |  Table 1: Weeks 3-5 ────────────────
t0 = doc.tables[0]
t1 = doc.tables[1]

# ─── Row 0-4 of Table 0: metadata rows ───────────────────────
def set_merged_row(table, row_idx, new_text):
    """Set text of a merged row (all cells share same text)."""
    row = table.rows[row_idx]
    for cell in row.cells:
        for p in cell.paragraphs:
            p.text = ""
        row.cells[0].paragraphs[0].text = new_text

# Row 0: Name
set_merged_row(t0, 0, "Name of Intern: Aryan Kori                                                        College: Barkatullah University, Bhopal")

# Row 1: Centre
set_merged_row(t0, 1, "Centre Name: Centre for Human Resource and Technology (CHR-T)                      Duration: 6th April 2026 - 5th June 2026")

# Row 2: Subject
set_merged_row(t0, 2, "Subject: Research and Policy Analysis")

# Row 3: Title
set_merged_row(t0, 3, "Title of the Study: Assessment of Digital Tool Adoption and Its Impact on Efficiency in MP Government Departments")

# Row 4: Overview
overview = (
    "Overview: This study evaluates how digital tools (e-Office, Bhulekh, RCMS, NREGASoft, ANMOL, etc.) "
    "are being adopted by government employees across four departments - Revenue, Rural Development, Forest, "
    "and Health - in Madhya Pradesh. Using the UTAUT framework and a bilingual survey instrument (English/Hindi), "
    "the research measures adoption rates, perceived efficiency gains, infrastructure gaps, and the role of "
    "personal tools (WhatsApp, AI) in official workflows. Target sample: 320 respondents across Class I-IV cadres."
)
set_merged_row(t0, 4, overview)

# ── Daily entries ─────────────────────────────────────────────
# Format: (sr, date, proposed_task, actual_task, attendance, remark)
# Week 1: 06-04 to 12-04
# Week 2: 13-04 to 19-04
# Week 3: 20-04 to 26-04
# Week 4: 27-04 to 03-05
# Week 5: 04-05 to 05-05

week1_entries = [
    ("1", "06.04.2026", "Conceptual understanding of the assigned topic",
     "Read a few government reports and circulars on how digital tools are being used in MP departments. Looked at e-Office and RCMS briefly to get a sense of what employees are expected to use. Made some rough notes.",
     "Present", ""),
    ("2", "07.04.2026", "Conceptual understanding of the assigned topic",
     "Continued going through reading material. Found some international reports on e-governance and compared where India stands. Tried to understand the bigger picture of why adoption matters at the state level.",
     "Present", ""),
    ("3", "08.04.2026", "Review of literature and understanding project/study objectives",
     "Went through research papers on technology adoption models. Read about a couple of frameworks that are commonly used in similar studies. Started getting a clearer idea of what the study should measure.",
     "Present", ""),
    ("4", "09.04.2026", "Identification of key themes and areas for literature review",
     "Figured out the main areas the study needs to cover - things like internet and device availability, how good the training is, what problems employees face, and whether they use personal apps for work. Noted which themes apply to which department.",
     "Present", ""),
    ("5", "10.04.2026", "Collection of relevant reports, papers, and secondary data",
     "Collected some useful documents - state IT policy papers, a few academic articles on e-governance in Indian states, and some portal usage data. Saved everything in one place for easy reference.",
     "Present", ""),
    ("6", "11.04.2026", "Identification of research gaps and key insights",
     "Went through everything collected so far. Most existing research looks at how citizens experience digital services, not how the employees feel about using them. That gap is basically what this study is about.",
     "Saturday", ""),
    ("7", "12.04.2026", "Sunday", "", "", ""),
]

week2_entries = [
    ("8", "13.04.2026", "Preparing the Concept Note / Proposal",
     "Started writing the concept note. Put together a title, short introduction, and the problem statement. Tried to keep it straightforward and professional.",
     "Present", ""),
    ("9", "14.04.2026", "Drafting background and rationale of the study",
     "Holiday but did some work from home. Wrote the background section - covered Digital India, MP's own e-governance push, and why it makes sense to study the employee side of things.",
     "General holiday", ""),
    ("10", "15.04.2026", "Defining objectives and scope of the study",
     "Wrote down the study objectives and defined the scope. Four departments, employees at different levels. Kept it focused so the fieldwork stays manageable.",
     "Present", ""),
    ("11", "16.04.2026", "Identifying key research questions / Drafting methodology approach",
     "Listed out the research questions and wrote a draft of the methodology. Planning to use a survey with rating scales plus some observations during field visits. Worked out the sampling plan.",
     "Present", ""),
    ("12", "17.04.2026", "Presentation of Concept Note / Proposal",
     "Presented the concept note to the mentor. Got feedback to be more specific about which portals each department uses. Also got a suggestion to ask about personal tools like WhatsApp that employees might be using for work.",
     "Present", ""),
    ("13", "18.04.2026", "Finalization of Concept Note / Proposal",
     "Made all the changes suggested by the mentor. Added the specific portal names for each department. Finalised the concept note and got the go-ahead to move forward.",
     "Saturday", ""),
    ("14", "19.04.2026", "Sunday", "", "", ""),
]

week3_entries = [
    ("15", "20.04.2026", "Developing the Research Framework",
     "Started working on the research framework. Read about a few existing models used in similar studies to understand how a framework is built. Began drafting a basic structure.",
     "Present", ""),
    ("16", "21.04.2026", "Understanding framework requirements and structure",
     "Figured out what the framework needs to show - how things like infrastructure, training, and ease of use connect to whether employees actually adopt digital tools. Drew a rough diagram on paper.",
     "Present", ""),
    ("17", "22.04.2026", "Identification of key variables and dimensions",
     "Listed the key variables for the study. On one side - device availability, internet quality, training, how easy the tools are, and whether seniors push for it. On the other side - actual adoption and efficiency. Each variable has a few survey questions linked to it.",
     "Present", ""),
    ("18", "23.04.2026", "Identification of key variables and dimensions",
     "Broke down the department-specific parts further. Each department has its own set of portals and issues - land records for Revenue, attendance apps for Rural Dev, GIS for Forest, patient tracking for Health. Drafted separate questions for each.",
     "Present", ""),
    ("19", "24.04.2026", "Validation with stakeholders/mentors",
     "Showed the draft framework to the mentor. Got approval on the overall structure. Was told to also add a section about personal tools since many employees seem to use WhatsApp and other apps to get work done.",
     "Present", ""),
    ("20", "25.04.2026", "Finalization of research framework",
     "Finalised the research framework after making all the changes. Prepared a neat diagram of the framework to be included in the proposal. Showed it to the mentor and got the go-ahead.",
     "Saturday", ""),
    ("21", "26.04.2026", "Sunday", "", "", ""),
]

week4_entries = [
    ("22", "27.04.2026", "Defining sub-themes and categorization based on framework",
     "Organised all the survey questions into proper sections - profile, infrastructure, performance, ease of use, social influence, awareness, training, challenges, and recommendations. Made sure each section matches a part of the framework.",
     "Present", ""),
    ("23", "28.04.2026", "Developing Key Performance Indicators (KPIs)",
     "Started making KPIs for each section. For infrastructure - things like how many employees have their own device and how they rate their internet. For training - how many actually got trained recently. Tried to keep indicators simple and measurable.",
     "Present", ""),
    ("24", "29.04.2026", "Developing Key Performance Indicators (KPIs)",
     "Continued making KPIs for the department-specific parts. Things like what percentage of land records still need paper, how often health workers enter the same data twice, whether forest staff have GPS devices. Put it all in one table.",
     "Present", ""),
    ("25", "30.04.2026", "Identifying required data points for each KPI",
     "Figured out what data will be needed for each KPI. Most of it comes from the survey itself. Some will need secondary data from department reports and government records.",
     "Present", ""),
    ("26", "01.05.2026", "Identifying required data points for each KPI",
     "Continued the same work for the remaining KPIs. Realised that a few things will need to be observed directly during field visits rather than just asked in the survey. Made a note of that.",
     "General holiday", ""),
    ("27", "02.05.2026", "Review and finalization of themes, KPIs, and data points",
     "Went through all the themes, KPIs and data points once more. Made a few small corrections and prepared the final clean version. Discussed it briefly with the mentor before wrapping up.",
     "Saturday", ""),
    ("28", "03.05.2026", "Sunday", "", "", ""),
]

week5_entries = [
    ("29", "04.05.2026", "Drafting data collection schedules/tools (questionnaires, formats)",
     "Started drafting the questionnaire. Divided the questions section by section. Made it bilingual - English and Hindi - so that employees at all levels can understand. Compiled the schedules as printable PDFs, one master and four department-wise.",
     "Present", ""),
    ("30", "05.05.2026", "Drafting data collection schedules/tools (questionnaires, formats)",
     "Finished the full draft of the questionnaire. Made it into a Google Form for easy online sharing. Also converted the PDFs to Word files. Organised all the files into proper folders and backed everything up.",
     "Present", ""),
]

def fill_row(table, row_idx, data):
    """Fill a single row with (sr, date, proposed, actual, attendance, remark)."""
    row = table.rows[row_idx]
    sr, date, proposed, actual, attendance, remark = data
    cells = row.cells
    # Clear existing text
    for cell in cells:
        for p in cell.paragraphs:
            p.text = ""
    cells[0].paragraphs[0].text = sr
    cells[1].paragraphs[0].text = date
    cells[2].paragraphs[0].text = proposed
    cells[3].paragraphs[0].text = actual
    cells[4].paragraphs[0].text = attendance
    cells[5].paragraphs[0].text = remark

# Table 0: Rows 5-6 are Week1 header + column header, Rows 7-13 are days 1-7
# Rows 14-15 are Week2 header + ... but wait, let me check again
# Row 5: Duration Week 1
# Row 6: Sr.No. / Date / Proposed / Actual / Attendance / Remark  (header)
# Row 7-13: Days 1-7
# Row 14: Duration Week 2
# Row 15-21: Days 8-14

# Fill Week 1
for i, entry in enumerate(week1_entries):
    fill_row(t0, 7 + i, entry)

# Fill Week 2
for i, entry in enumerate(week2_entries):
    fill_row(t0, 15 + i, entry)

# Table 1: 
# Row 0: Duration Week 3
# Row 1: header row
# Row 2-8: Days 15-21
# Row 9: Duration Week 4
# Row 10-16: Days 22-28
# Row 17: Duration Week 5
# Row 18-19: Days 29-30

for i, entry in enumerate(week3_entries):
    fill_row(t1, 2 + i, entry)

for i, entry in enumerate(week4_entries):
    fill_row(t1, 10 + i, entry)

for i, entry in enumerate(week5_entries):
    fill_row(t1, 18 + i, entry)

doc.save(dst)
print(f"Saved to: {dst}")
