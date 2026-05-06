from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc_path = r"c:\Users\aryan\Downloads\Framework_Explanation.docx"

doc = Document()

# Title
title = doc.add_heading("Framework for Interns: Simple Explanation", level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Introduction
doc.add_paragraph("This document explains the items placed in your Research Framework Excel sheet. The framework is built around your 4 core Research Objectives (RO). Each objective acts as a guide, telling us why we are doing the research and what we want to achieve. The questions and indicators then break down those objectives into simple, actionable steps.")

doc.add_heading("Objective 1: Evaluate Current Wellness Status", level=1)
p = doc.add_paragraph()
p.add_run("Why we have this objective: ").bold = True
p.add_run("To understand the baseline health and wellness of the employees before making any big changes. We need to know where they stand right now.")
p = doc.add_paragraph()
p.add_run("What we will do with it: ").bold = True
p.add_run("We will ask straightforward questions about their physical health, stress levels, and workspace comfort. This gives us a clear picture of the current situation.")
p = doc.add_paragraph()
p.add_run("How the questions define this: ").bold = True
p.add_run("We ask things like, \"How often do you exercise?\", \"Rate your stress level\", and \"Are you satisfied with your office lighting?\" – these direct questions map directly back to understanding their current wellness status.")

doc.add_heading("Objective 2: Understand How Absence of Wellness Affects Performance", level=1)
p = doc.add_paragraph()
p.add_run("Why we have this objective: ").bold = True
p.add_run("To connect the dots between feeling unwell and working inefficiently. We need to prove that health impacts daily work.")
p = doc.add_paragraph()
p.add_run("What we will do with it: ").bold = True
p.add_run("We will look at sick leaves, job satisfaction, and how financial or personal stress stops employees from focusing.")
p = doc.add_paragraph()
p.add_run("How the questions define this: ").bold = True
p.add_run("Questions like \"How many sick leaves have you taken?\" and \"Does financial stress affect your focus?\" directly measure the impact of wellness on performance.")

doc.add_heading("Objective 3: Explore How Age, Gender, and Roles Shape Wellness Needs", level=1)
p = doc.add_paragraph()
p.add_run("Why we have this objective: ").bold = True
p.add_run("Because a 50-year-old senior officer has different stress points than a 25-year-old new recruit. We need to ensure support isn't \"one size fits all.\"")
p = doc.add_paragraph()
p.add_run("What we will do with it: ").bold = True
p.add_run("We will group the answers by age, gender, and job class to see the unique challenges each group faces.")
p = doc.add_paragraph()
p.add_run("How the questions define this: ").bold = True
p.add_run("We ask for demographic details (Age, Gender, Class) and qualitative questions like \"Describe the most stressful part of your workday.\" This helps us tailor specific wellness solutions.")

doc.add_heading("Objective 4: Develop an Evidence-Based Framework for Increased Efficiency", level=1)
p = doc.add_paragraph()
p.add_run("Why we have this objective: ").bold = True
p.add_run("To create a final, practical recommendation for the government. We want to show that investing in employee wellness leads to better public service.")
p = doc.add_paragraph()
p.add_run("What we will do with it: ").bold = True
p.add_run("We will combine all the data to create a \"Wellness Index\" and see how existing schemes (like AYUSH or ABHA) are currently being used, identifying gaps we can fix.")
p = doc.add_paragraph()
p.add_run("How the questions define this: ").bold = True
p.add_run("Questions asking about their awareness of current government initiatives and overall self-rated efficiency help us build this final, actionable framework.")

doc.add_heading("Understanding the Excel Columns", level=1)
doc.add_paragraph("Indicators: The specific things we are looking for (like sleep quality, stress level, sick leaves).", style='List Bullet')
doc.add_paragraph("Quantitative/Qualitative: How we measure it. Quantitative means numbers (surveys). Qualitative means stories and descriptions (interviews).", style='List Bullet')
doc.add_paragraph("Analysis Tools & Visuals: How we will make sense of the data (using simple charts like Pie charts, Bar charts, and basic statistics).", style='List Bullet')
doc.add_paragraph("Levels: Who we are asking (Class I to IV employees).", style='List Bullet')
doc.add_paragraph("Means of Verification: Proof of our data (like the survey answers, HR leave records, or interview transcripts).", style='List Bullet')

doc.save(doc_path)
print(f"Explanation document saved to: {doc_path}")
