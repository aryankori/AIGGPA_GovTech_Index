"""
Generate a Word document: Background section for the AIGGPA IT Utilisation study.
Written in plain, simple language as requested.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Default font ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.5

# ── Heading styles ────────────────────────────────────────────
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x00, 0x35, 0x66)  # navy

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)


def add_para(text, bold=False, italic=False, indent=False):
    """Helper to add a normal paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_spaced_heading(text, level=1):
    """Add a heading with a little space above."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(18)
    return h


# ══════════════════════════════════════════════════════════════
#  TITLE
# ══════════════════════════════════════════════════════════════
title = doc.add_heading('Background of the Study', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Assessment of IT Tool Utilisation Among Government Officials\n'
                 'School Education, Health & Forest Departments — Government of Madhya Pradesh')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r.italic = True

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════
#  1. WHY ARE WE DOING THIS STUDY?
# ══════════════════════════════════════════════════════════════
add_spaced_heading('1.  Why Are We Doing This Study?', level=2)

add_para(
    'The Government of Madhya Pradesh has spent years building digital tools '
    'for its departments. Agencies like MAP-IT and NIC have created websites, '
    'portals, and mobile apps so that government officers can do their work '
    'faster, with less paper, and with better data. On paper, these tools '
    'exist and are ready to use.'
)

add_para(
    'But here is the problem: most officers at the district and block level '
    'still do not use them properly. Instead of logging into the official portal, '
    'they keep using old methods — physical registers, paper files sent by post, '
    'WhatsApp groups, and personal Excel sheets. The fancy digital systems sit '
    'there, largely unused.'
)

add_para(
    'Nobody has properly studied why this is happening in Madhya Pradesh. '
    'We hear complaints in training rooms at AIGGPA (Atal Bihari Vajpayee '
    'Institute of Good Governance and Policy Analysis, Bhopal) all the time: '
    '"We know the portal exists, but we don\'t use it." But no one has sat '
    'down to systematically figure out which tools are being ignored, and '
    'exactly what is stopping officers from using them.'
)

add_para(
    'That is what this study does. It is a ground-level investigation into '
    'three major departments — School Education, Health, and Forest — to '
    'find out which IT tools are available, which ones are actually being '
    'used, and what the real barriers are. The end goal is to give AIGGPA '
    'solid evidence so it can design better training programmes that actually '
    'solve the real problems.'
)

# ══════════════════════════════════════════════════════════════
#  2. WHAT DOES EXISTING RESEARCH SAY? (Literature Review)
# ══════════════════════════════════════════════════════════════
add_spaced_heading('2.  What Does Existing Research Tell Us?  (Literature Review)', level=2)

add_para(
    'Researchers around the world have studied why government employees '
    'struggle with technology. Here is what the key studies say, explained '
    'simply:'
)

# -- Heeks
add_spaced_heading('The "Design-Reality Gap" (Heeks, 2003)', level=3)
add_para(
    'Richard Heeks, a researcher at the University of Manchester, found that '
    'most government IT projects in developing countries fail — not because '
    'the software is bad, but because the people who build the system and '
    'the people who use it live in completely different worlds. The developers '
    'assume that offices have good internet, that staff know how to use '
    'computers, and that bosses will encourage digital work. In reality, '
    'none of that may be true. Heeks calls this the "design-reality gap." '
    'The bigger the gap, the more likely the system will just sit unused.'
)

# -- Bhatnagar
add_spaced_heading('Barriers to IT in Public Administration (Bhatnagar, 2004)', level=3)
add_para(
    'Subhash Bhatnagar studied IT adoption in Indian government offices and '
    'found two recurring problems: (1) there are no official rules that say '
    '"you must use the portal" — so officers treat it as optional, and '
    '(2) front-line staff rarely get proper computer training. Without '
    'clear rules and without training, digital tools become decoration, '
    'not infrastructure.'
)

# -- UTAUT
add_spaced_heading('UTAUT — A Framework to Understand Technology Adoption (Venkatesh et al., 2003)', level=3)
add_para(
    'Venkatesh and his colleagues created a model called UTAUT (Unified '
    'Theory of Acceptance and Use of Technology). It is the most widely '
    'used framework for understanding why people accept or reject '
    'technology in organisations. It boils everything down to four factors:'
)

add_para('Performance Expectancy (PE) — "Will this tool actually help me do my job better?"', bold=True, indent=True)
add_para('Effort Expectancy (EE) — "Is this tool easy enough for me to use?"', bold=True, indent=True)
add_para('Social Influence (SI) — "Is my boss telling me to use it? Are my colleagues using it?"', bold=True, indent=True)
add_para('Facilitating Conditions (FC) — "Do I even have a working computer and internet to use it?"', bold=True, indent=True)

add_para(
    'UTAUT is especially useful for government settings because it captures '
    'something that purely psychological models miss: the role of hierarchy. '
    'In a government office, if a senior officer insists on a physical '
    'signature on paper, it does not matter how easy the digital system is — '
    'the subordinate has no choice but to use paper. UTAUT accounts for '
    'this through Social Influence.'
)

# -- Legacy systems / Global context
add_spaced_heading('The Legacy System Problem (GAO, 2019; Cordella & Bonina, 2012)', level=3)
add_para(
    'This is not just an Indian problem. A 2019 report by the U.S. Government '
    'Accountability Office (GAO) found that government agencies worldwide spend '
    'up to 80% of their IT budgets just keeping old, outdated systems running — '
    'leaving almost nothing for building modern tools. Cordella and Bonina (2012) '
    'argued that real digital reform in government is not just about buying new '
    'software; it is about changing the entire way public services are delivered, '
    'which requires political will, training, and structural changes.'
)

# -- Digital India
add_spaced_heading('Digital India and Mission Karmayogi', level=3)
add_para(
    'The Indian government recognises this challenge at the national level. '
    'The Digital India programme (2015) was launched to push every government '
    'service online. More recently, Mission Karmayogi (2020) — officially called '
    'the National Programme for Civil Services Capacity Building — shifted the '
    'approach to government training. Instead of generic classroom sessions, '
    'it now wants training that is tied to a specific officer\'s role and '
    'the specific digital tools they need to use. The iGOT Karmayogi platform '
    'is the online platform where these role-specific courses are hosted.'
)

add_para(
    'AIGGPA, as Madhya Pradesh\'s top training institution for government '
    'officers, is required to design and deliver these iGOT modules. But to '
    'design good training, you first need to know what the actual problems are. '
    'That is exactly the gap this study fills.'
)

# ══════════════════════════════════════════════════════════════
#  3. WHAT IT TOOLS EXIST TODAY?
# ══════════════════════════════════════════════════════════════
add_spaced_heading('3.  What IT Tools Already Exist in These Departments?', level=2)

add_para(
    'Before we can study whether tools are being used, it helps to understand '
    'what is actually available. Here is a quick summary of what the three '
    'departments have:'
)

add_spaced_heading('School Education Department', level=3)
add_para(
    'Education Portal 3.0 — This is a massive system that manages roughly '
    '3,50,000 permanent teachers across more than 92,000 schools in Madhya Pradesh. '
    'It handles teacher transfers, promotions, attendance tracking, and academic '
    'monitoring. Despite its scale, many district and block-level officers do not '
    'use it consistently.'
)

add_spaced_heading('Health Department', level=3)
add_para(
    'The Health Department runs HRMIS (Human Resource Management Information System) '
    'for managing its staff. It also operates NHM CHETNA, a platform used to train '
    'front-line health workers like ANMs and ASHAs. Recently, the department launched '
    'an AI chatbot called "Ayushman Sakhi" that helps citizens find hospitals and '
    'access health information.'
)

add_spaced_heading('Forest Department', level=3)
add_para(
    'The Forest Department uses satellite data to generate real-time alerts about '
    'forest encroachments. All wildlife tourism permits are processed through the '
    'MPOnline portal.'
)

add_para(
    'In short, the tools exist. The technology is there. The question is: why '
    'aren\'t the officers using them?',
    italic=True
)

# ══════════════════════════════════════════════════════════════
#  4. EFFECTIVENESS OF IT TOOLS & WHY THIS STUDY MATTERS
# ══════════════════════════════════════════════════════════════
add_spaced_heading('4.  How Effective Are IT Tools — and Why Does This Study Matter?', level=2)

add_para(
    'When government IT tools are actually used properly, the benefits are '
    'enormous:'
)

benefits = [
    ('Faster work:', 'Tasks that take days on paper — like processing a teacher transfer or '
     'tracking health worker attendance — can be done in minutes on a portal.'),
    ('Better data:', 'Instead of guessing how many schools reported attendance this month, '
     'a state official can pull up a real-time dashboard. Policy decisions become '
     'data-driven instead of gut-driven.'),
    ('Less corruption and manipulation:', 'When files move through a digital system, '
     'there is a clear audit trail. It becomes much harder for someone to "sit on a file" '
     'or lose it deliberately.'),
    ('Money saved:', 'Digital processing cuts down on printing, courier costs, and the '
     'need for physical storage of thousands of registers.'),
]

for label, desc in benefits:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    b = p.add_run(label + ' ')
    b.bold = True
    p.add_run(desc)

add_para(
    'But all of these benefits remain theoretical if the tools are not being used. '
    'And right now, they are not — at least not uniformly. That is the core problem.'
)

add_spaced_heading('Why No One Has Done This Before', level=3)
add_para(
    'Madhya Pradesh has never had a systematic, evidence-based study that maps '
    'out which IT tools are being under-used and why, department by department, '
    'level by level. There are anecdotal complaints — trainers hear them, '
    'senior officers suspect them — but no one has collected structured data. '
    'Without that data, training programmes are designed on assumptions, not '
    'evidence. This study provides that missing evidence.'
)

add_spaced_heading('What This Study Will Produce', level=3)
add_para(
    'Using the UTAUT framework, every barrier found in this study will be '
    'classified into one of the four categories (PE, EE, SI, FC). '
    'This classification directly tells AIGGPA what kind of intervention '
    'is needed:'
)

interventions = [
    ('If the barrier is PE or EE (the officer doesn\'t see value or finds it '
     'hard to use):', ' → AIGGPA designs a targeted training module for the iGOT platform.'),
    ('If the barrier is SI or FC (the boss discourages it, or there\'s no '
     'working computer):', ' → AIGGPA sends a policy brief to the Department Secretary '
     'recommending structural fixes — like mandating digital submissions or '
     'funding hardware repairs.'),
]

for label, desc in interventions:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    b = p.add_run(label)
    b.bold = True
    p.add_run(desc)

add_para(
    'In other words, this study does not just say "officers need training." '
    'It figures out exactly what kind of help each group of officers needs, '
    'and whether that help is training, better equipment, or a policy change '
    'from above.'
)

# ══════════════════════════════════════════════════════════════
#  5. THE NEED FOR THIS STUDY — SUMMING IT UP
# ══════════════════════════════════════════════════════════════
add_spaced_heading('5.  The Need for This Study — A Summary', level=2)

add_para(
    'To put it all together in the simplest possible terms:'
)

needs = [
    'The MP government has built IT tools for its departments.',
    'Officers are not using these tools properly — they default to paper, '
    'WhatsApp, and personal workarounds.',
    'Nobody has systematically studied why this is happening in MP.',
    'Existing research (Heeks, Bhatnagar, Venkatesh) tells us the likely reasons: '
    'bad design fit, lack of rules, no training, boss pressure, and missing '
    'hardware — but we need local data to confirm which of these apply here.',
    'AIGGPA is responsible for training government officers (under Mission '
    'Karmayogi / iGOT). To design good training, it needs hard evidence about '
    'real gaps.',
    'This study collects that evidence by interviewing 54 officers across three '
    'departments, observing their workstations, and getting IT usage statistics.',
    'The findings will be classified using the UTAUT framework and turned into '
    'actionable recommendations — some for training design, some for policy '
    'changes.',
]

for i, item in enumerate(needs, 1):
    p = doc.add_paragraph()
    p.style = doc.styles['List Number']
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{item}')

# ══════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════
add_spaced_heading('References', level=2)

refs = [
    'Bhatnagar, S. (2004). IT Platforms: From Vision to Implementation: '
    'A Practical Guide with Case Studies. Sage Publications.',

    'Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. '
    'Qualitative Research in Psychology, 3(2), 77–101.',

    'Cordella, A., & Bonina, C. M. (2012). A public value perspective for '
    'ICT enabled public sector reforms. Government Information Quarterly, '
    '29(4), 512–520.',

    'Government of India, Department of Personnel & Training. (2020). '
    'National Programme for Civil Services Capacity Building (Mission Karmayogi).',

    'Government of Madhya Pradesh, School Education Department. (2025). '
    'Education Portal 3.0. https://sederp.educationportal3.in',

    'Heeks, R. (2003). Most IT-for-Development Projects Fail: How Can Risks '
    'Be Reduced? iGovernment Working Paper No. 14, University of Manchester.',

    'Ministry of Electronics & Information Technology, Government of India. '
    '(2015). Digital India Programme. https://digitalindia.gov.in/',

    'United States Government Accountability Office (GAO). (2019). '
    'Agencies Need to Develop Modernization Plans for Critical Legacy Systems '
    '(GAO-19-471).',

    'Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). '
    'User acceptance of information technology: Toward a unified view. '
    'MIS Quarterly, 425–478.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.size = Pt(10)

# ── Save ──────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'Background_of_the_Study.docx')
doc.save(output_path)
print(f'✅  Document saved to: {output_path}')
