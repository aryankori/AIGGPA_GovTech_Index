import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_latex(text):
    text = text.replace(r'$\square$', '☐')
    text = text.replace(r'\likert', '☐ 1   ☐ 2   ☐ 3   ☐ 4   ☐ 5')
    text = text.replace(r'\yesno', '☐ Yes   ☐ No')
    text = text.replace(r'\blank', '_' * 25)
    text = re.sub(r'\\rule\{.*?\}\{.*?\}', '_' * 25, text)
    text = re.sub(r'\{\\scriptsize\((.*?)\)\}', r'(\1)', text)
    text = re.sub(r'\\[a-zA-Z]+\s*', '', text) # remove remaining simple commands
    text = text.replace('{', '').replace('}', '')
    text = text.replace(r'\\', '')
    text = text.replace('[6pt]', '').replace('[5pt]', '').replace('[4pt]', '').replace('[3pt]', '').replace('[2pt]', '')
    text = text.strip()
    return text

def parse_and_convert(tex_path, docx_path):
    if not os.path.exists(tex_path):
        return

    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    with open(tex_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    current_table = None

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        
        if not line or line.startswith('%') or line.startswith(r'\documentclass') or line.startswith(r'\usepackage') or line.startswith(r'\pagestyle') or line.startswith(r'\newcolumntype') or line.startswith(r'\newcommand'):
            continue
            
        if r'\begin{document}' in line:
            continue
        if r'\end{document}' in line:
            break
            
        # Headers
        if r'{\Large\bfseries' in line or r'{\large\bfseries' in line:
            header_text = re.search(r'bfseries(.*?)}', line)
            if header_text:
                p = doc.add_paragraph()
                run = p.add_run(clean_latex(header_text.group(1)))
                run.bold = True
                run.font.size = Pt(14 if 'Large' in line else 12)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if 'center' in line else WD_ALIGN_PARAGRAPH.LEFT
            continue

        if r'\noindent\textbf{' in line and not in_table:
            text = re.search(r'textbf{(.*?)}', line)
            if text:
                p = doc.add_paragraph()
                run = p.add_run(clean_latex(text.group(1)) + " ")
                run.bold = True
                rest = line[line.find('}')+1:]
                if rest:
                    p.add_run(clean_latex(rest))
            continue
            
        # Tables
        if r'\begin{tabularx}' in line:
            in_table = True
            current_table = doc.add_table(rows=0, cols=3)
            current_table.autofit = False
            # Set column widths roughly
            for cell in current_table.columns[0].cells: cell.width = Inches(0.5)
            for cell in current_table.columns[1].cells: cell.width = Inches(3.5)
            for cell in current_table.columns[2].cells: cell.width = Inches(3.0)
            continue
            
        if r'\end{tabularx}' in line:
            in_table = False
            doc.add_paragraph() # spacing
            continue
            
        if in_table:
            # handle table rows
            # might span multiple lines if there's no \\ at the end, but in our tex they usually have \\ at the end.
            if '&' in line:
                parts = line.split('&')
                row = current_table.add_row()
                for col_idx in range(min(3, len(parts))):
                    cell = row.cells[col_idx]
                    cell.text = clean_latex(parts[col_idx])
            continue
            
        # Normal text or other stuff
        cleaned = clean_latex(line)
        if cleaned:
            doc.add_paragraph(cleaned)

    doc.save(docx_path)
    print(f"Saved {docx_path}")

files_to_convert = [
    ("Schedule_Revenue.tex", "Schedule_Revenue.docx"),
    ("Schedule_Rural_Development.tex", "Schedule_Rural_Development.docx"),
    ("Schedule_Forest.tex", "Schedule_Forest.docx"),
    ("Schedule_Health.tex", "Schedule_Health.docx"),
    ("AIGGPA_Printable_Schedule.tex", "AIGGPA_Printable_Schedule.docx")
]

downloads_dir = r"c:\Users\aryan\Downloads"

for tex_file, docx_file in files_to_convert:
    tex_path = os.path.join(os.getcwd(), tex_file)
    docx_path = os.path.join(downloads_dir, docx_file)
    try:
        parse_and_convert(tex_path, docx_path)
    except Exception as e:
        print(f"Error converting {tex_file}: {e}")
