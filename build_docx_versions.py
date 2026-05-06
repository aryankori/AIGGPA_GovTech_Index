"""
Builds DOCX versions of Final Proposal and Methodology Scope Options.
Retains exact wordings (no em dashes, uses 'government employees').
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(__file__)

def setup_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    return doc

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def build_final_proposal():
    doc = setup_document()
    
    # Title Page
    doc.add_paragraph()
    doc.add_paragraph()
    title = add_heading(doc, 'DIGITAL ASSESSMENT OF GOVERNMENT EMPLOYEES', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.bold = True
    
    sub = doc.add_paragraph('Research Proposal')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph('\n\n\n\n')
    info = doc.add_paragraph()
    info.add_run('Prepared By:\n').bold = True
    info.add_run('Aryan Kori\n\n')
    info.add_run('Institution:\n').bold = True
    info.add_run('AIGGPA, Bhopal\n\n')
    info.add_run('Departments:\n').bold = True
    info.add_run('School Education, Health, Forest\n(Govt. of Madhya Pradesh)\n\n')
    info.add_run('Date:\n').bold = True
    info.add_run('April 2026')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Content
    add_heading(doc, '1. INTRODUCTION AND THE DIGITAL GAP')
    intro = (
        "India is moving from paper records to digital systems to improve public offices. "
        "The use of digital technology in public administration is no longer optional. It has become essential for improving efficiency, transparency, and service delivery. Across the world, governments are adopting digital systems, but their success depends on how well government employees are able to understand and use these technologies in their daily work. "
        "While digital infrastructure in government has expanded, there is still a gap in awareness, usage, and skill levels among employees. While modern organizations adopt digital tools quickly and use them to improve speed and efficiency, government departments often face challenges such as limited training, lack of technical support, and resistance to change. Studies by McKinsey & Company (2018), specifically their report 'Unlocking Success in Digital Transformations', show that even with high investment in IT systems, actual usage at the employee level remains low, which reduces the expected benefits. "
        "In Madhya Pradesh, departments such as Revenue, Rural Development, Forest, and Health play a key role in grassroots service delivery. Because these departments manage vast geographical areas and massive rural populations, their digital transition is highly complex but absolutely critical. While they rely on digital systems for functions like land records, rural schemes, health monitoring, and forest management, the scale of operations makes implementation difficult. It is important to assess the level of digital infrastructure, awareness among government employees, and the ongoing training support provided by these departments to understand the operational gaps preventing full digital adoption. "
        "The purpose of this study is to identify these gaps and suggest practical steps for improving the use of digital technology by government employees in Madhya Pradesh. The focus is on strengthening employee capacity, reducing delays, and improving the overall quality and transparency of public service delivery."
    )
    doc.add_paragraph(intro)

    add_heading(doc, '2. RESEARCH OBJECTIVES')
    objectives = [
        "To know about the digital infrastructure, awareness and use of various tools and technologies by government employees in their office work.",
        "To assess the capacity building efforts made by the department on technological support to their government employees for using digital tools.",
        "To identify bottlenecks, issues and challenges faced by government employees using technology or digital tools.",
        "To suggest the required actions for the department for efficient delivery service by using appropriate digital tools and technologies."
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Number')

    add_heading(doc, '3. HYPOTHESIS')
    hypotheses = [
        "Government employees who know about digital tools are more likely to use them.",
        "Poor training and old computers are the main reasons government employees avoid digital tools.",
        "Government employees use digital tools more when their seniors tell them to.",
        "Users of digital tools believe they work better than those using paper.",
        "Offices with better IT support have more people using digital tools."
    ]
    for hyp in hypotheses:
        doc.add_paragraph(hyp, style='List Number')

    add_heading(doc, '4. IMPORTANT QUESTIONS THIS STUDY WILL ANSWER')
    questions = [
        "What is the current status of digital infrastructure in the selected departments, and to what extent are government employees aware of and using various digital tools and technologies in their daily office work?",
        "What kind of capacity building efforts, training programs, and technical support are being provided by the departments to help government employees use digital tools effectively?",
        "What are the major bottlenecks, issues, and challenges faced by government employees while using digital tools and technologies in their work?",
        "What actions and measures can be suggested to improve the use of digital tools for more efficient and effective service delivery in government departments?"
    ]
    for q in questions:
        doc.add_paragraph(q, style='List Number')

    add_heading(doc, '5. FRAMEWORK')
    doc.add_paragraph(
        "This study uses a standard model called UTAUT to understand why people use technology. "
        "It looks at four areas: if the tool helps work, if it is easy to use, if coworkers use it, "
        "and if the right equipment is available. This model helps us see if problems are personal "
        "or organizational."
    )
    doc.add_paragraph("The study follows this flow: Available Tools > Factors Affecting Use > Level of Use > Impact on Work.")
    utaut_items = [
        "Expected Benefit: Does the government employee think the tool makes work faster?",
        "Ease of Use: Is the system harder than using paper?",
        "Social Pressure: Do bosses and coworkers encourage usage?",
        "Support: Is there a working computer, internet, and tech help?"
    ]
    for item in utaut_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '6. SCOPE')
    doc.add_paragraph("This study focuses on how government employees in Madhya Pradesh use IT tools for internal work.")
    scope_items = [
        "Location: Offices in Madhya Pradesh, covering selected districts or divisions.",
        "Participants: Government employees from the Revenue, Rural Development, Forest, and Health departments across administrative levels.",
        "Focus: Internal office tools only, not public services.",
        "Duration: 2 to 3 months."
    ]
    for scope in scope_items:
        doc.add_paragraph(scope, style='List Bullet')

    add_heading(doc, '7. METHODOLOGY')
    doc.add_paragraph("The study will use two types of data:")
    add_heading(doc, 'Sampling Strategy', level=2)
    samp_items = [
        "Stratified Random Sampling: Used for selecting government employees to ensure equal representation from all four designated departments. The sample will be further stratified based on designation (senior, mid-level, and junior administrative roles) and location (headquarters versus district-level offices) to capture diverse operational challenges.",
        "Purposive Sampling: Used exclusively for selecting key informants, such as departmental IT officers, system administrators, and senior policy officials. These individuals are selected because they possess specialized knowledge of backend infrastructure and institutional bottlenecks that frontline government employees may not have."
    ]
    for st in samp_items:
        doc.add_paragraph(st, style='List Bullet')
    add_heading(doc, 'Data Collection and Analysis', level=2)
    doc.add_paragraph("Interviews with government employees and IT personnel to understand their daily challenges. We will also check office equipment in person.")
    meth_tools = [
        "Interview guides for government employees and IT personnel.",
        "Checklists to record details of computer and internet availability."
    ]
    for tool in meth_tools:
        doc.add_paragraph(tool, style='List Bullet')
    doc.add_paragraph("Quantitative data will be analyzed using tools such as Microsoft Excel and SPSS to identify patterns, trends, and relationships between digital technology usage and work efficiency. Comparative analysis will be carried out across the four different government departments to identify which sectors are adapting faster. Qualitative data will be analyzed using thematic analysis to identify key issues, bottlenecks, and areas for improvement among government employees.")

    add_heading(doc, '8. TIMELINE')
    doc.add_paragraph("The total project will take 8 weeks:")
    timeline = [
        "Weeks 1-2: Preparation and design.",
        "Weeks 3-5: Visits and interviews.",
        "Weeks 6-7: Analysis of findings.",
        "Week 8: Final report and presentation."
    ]
    for t in timeline:
        doc.add_paragraph(t, style='List Bullet')

    add_heading(doc, '9. BUDGET')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Amount (Rs.)'
    budget = [
        ("Design and Literature Review", "32,000"),
        ("Data Collection and Travel", "15,500"),
        ("Software and Analysis", "8,000"),
        ("Writing and Printing", "4,500"),
        ("TOTAL", "60,000")
    ]
    for cat, amt in budget:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = amt

    add_heading(doc, '10. EXPECTED CONTRIBUTION')
    doc.add_paragraph(
        "This study will show the real gap between available IT tools and actual use. "
        "It will identify specific problems and suggest practical steps to help government employees "
        "work digitally. The findings will help other offices improve their efficiency."
    )

    add_heading(doc, '11. LIMITATIONS')
    doc.add_paragraph(
        "The study is limited to Bhopal and nearby areas. With a sample of 60 people, it gives "
        "a good overview of challenges but may not represent every office in the state. We rely on "
        "honest answers from government employees, but in-person checks will help confirm the details."
    )

    add_heading(doc, '12. ACRONYMS AND KEY TERMS')
    acronyms = [
        "AIGGPA: Atal Bihari Vajpayee Institute of Good Governance and Policy Analysis.",
        "Govt. of MP: Government of Madhya Pradesh.",
        "UTAUT (Unified Theory of Acceptance and Use of Technology): A standard research model used to understand why people choose to use new technologies, by looking at factors like how helpful or easy the tool is.",
        "Digital Gap: The difference between available digital tools (like websites and apps) and their actual use by employees.",
        "Primary Data: New information gathered directly from interviews and observations during this study.",
        "Secondary Data: Existing information and usage records reviewed from IT departments.",
        "Thematic analysis: A method for grouping interview notes and findings into themes to make sense of the results."
    ]
    for acr in acronyms:
        doc.add_paragraph(acr, style='List Bullet')

    add_heading(doc, '13. REFERENCES')
    refs = [
        'Alrawabdeh, W. (2014). Examining factors affecting digital platform adoption in Jordan. International Journal of Business and Social Science, 5(8), 232-241.',
        'Bhatnagar, S. (2004). E-Government: From Vision to Implementation - A Practical Guide with Case Studies. Sage Publications.',
        'Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77-101.',
        'Cordella, A., & Bonina, C. M. (2012). A public value perspective for ICT enabled public sector reforms. Government Information Quarterly, 29(4), 512-520.',
        'Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.',
        'Government of India, Department of Personnel & Training. (2020). National Programme for Civil Services Capacity Building (Mission Karmayogi).',
        'Government of India, Ministry of Electronics & Information Technology. (2015). Digital India. https://digitalindia.gov.in/',
        'Heeks, R. (2003). Most eGovernment-for-Development Projects Fail: How Can Risks Be Reduced? University of Manchester.',
        'Heeks, R. (2006). Implementing and Managing eGovernment. Sage Publications.',
        'Kumar, R., & Best, M. L. (2006). Impact and sustainability of e-government services in developing countries. The Information Society, 22(1), 1-12.',
        'Ndou, V. (2004). E-government for developing countries. The Electronic Journal of Information Systems in Developing Countries, 18(1), 1-24.',
        'OECD. (2014). Recommendation of the Council on Digital Government Strategies.',
        'United Nations. (2022). E-Government Survey 2022: The Future of Digital Government.',
        'United States GAO. (2019). Agencies Need to Develop Modernization Plans for Critical Legacy Systems (GAO-19-471).',
        'Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425-478.',
        'Venkatesh, V., Thong, J. Y. L., & Xu, X. (2012). Consumer acceptance and use of information technology. MIS Quarterly, 36(1), 157-178.',
        'World Bank. (2016). World Development Report 2016: Digital Dividends.'
    ]
    for r in refs:
        doc.add_paragraph(r)

    doc.save(os.path.join(BASE, 'Final_Proposal_Formatted.docx'))
    print("Saved Final_Proposal_Formatted.docx")

def build_methodology_options():
    doc = setup_document()
    
    # Title Page
    doc.add_paragraph('\n\n\n')
    title = add_heading(doc, 'SCOPE AND METHODOLOGY OPTIONS PROPOSAL', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.bold = True
    
    sub = doc.add_paragraph('Digital Assessment of Government Employees')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

    # Content
    add_heading(doc, 'PURPOSE OF THIS DOCUMENT')
    doc.add_paragraph(
        "This document outlines three different levels of scope for the research project, ranging "
        "from a simple, fast approach to a highly complex, algorithmic approach. It details the "
        "Tools, Processes, Data Collection Methods, Sampling Techniques, and Sorting/Analysis "
        "Algorithms for each level.\n\n"
        "This will help management decide the final commitment level, required resources, "
        "and timeline for the study."
    )

    opts = doc.add_paragraph()
    opts.add_run('OPTION 1: BASIC SCOPE (Qualitative Focus)\n').bold = True
    opts.add_run('Fast, low cost, interview-based. (Current Proposal Level)\n\n')
    opts.add_run('OPTION 2: MODERATE SCOPE (Quantitative Focus)\n').bold = True
    opts.add_run('Survey-based, statistically significant, uses basic algorithms.\n\n')
    opts.add_run('OPTION 3: COMPLEX SCOPE (Mixed-Methods + ML)\n').bold = True
    opts.add_run('Massive scale, uses Structural Equation Modeling and Clustering.')

    doc.add_page_break()

    add_heading(doc, 'OPTION 1: BASIC SCOPE')
    details1 = [
        ("Summary: ", "This is the simplest approach. It focuses on talking to a small group of people to understand their problems deeply, rather than collecting numbers. This is the baseline proposed earlier."),
        ("Sampling Method: ", "Purposive & Convenience Sampling. We carefully select 45 to 60 government employees who represent different roles (clerks, managers, IT staff) in Bhopal. We do not use randomized formulas."),
        ("Collection Tools: ", "In-person, one-on-one structured interviews (30 minutes each). Notes and audio recordings are collected. A checklist is used to observe their physical workspace (computer condition)."),
        ("Sorting & Analysis Algorithm: ", "Thematic Analysis (Manual Coding). Data is transcribed into Excel/Word. We manually read the answers and sort them into four 'buckets' or themes based on the UTAUT model (e.g., all answers about bad internet go into the 'Facilitating Conditions' bucket)."),
        ("Time & Commitment: ", "8 Weeks. Low software cost. Requires 1 researcher visiting offices heavily for 3 weeks."),
        ("Pros: ", "Very fast. Gives real, human stories and specific software complaints. Very low cost."),
        ("Cons: ", "Sample size is too small to say 'X percent of the whole state feels this way'. It is an exploratory study.")
    ]
    for title, text in details1:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(text)

    doc.add_paragraph('\n')
    add_heading(doc, 'OPTION 2: MODERATE SCOPE')
    details2 = [
        ("Summary: ", "This approach uses numbers to prove a point. We ask hundreds of government employees to rate their digital experience on a 1 to 5 scale. This allows us to create graphs and calculate averages."),
        ("Sampling Method: ", "Stratified Random Sampling. We take the total population of government employees in 3 departments, divide them into groups (strata) based on their rank (e.g., senior, mid-level, junior), and randomly select 300 to 400 people across Bhopal and 2 other cities."),
        ("Collection Tools: ", "Online Likert-Scale Surveys (Google Forms or Qualtrics) sent via official email or WhatsApp groups. Paper surveys act as backups for those without internet. The survey asks them to agree/disagree with 20 statements."),
        ("Sorting & Analysis Algorithm: ", "Descriptive Statistics (SPSS or Python Pandas/SciPy). We use sorting algorithms to group data by age, department, and rank. We use algorithms like T-tests and ANOVA to check if there is a real mathematical difference between how young vs. older employees use tools."),
        ("Time & Commitment: ", "12 to 14 Weeks. Requires management to push out the survey link to hundreds of people. Requires basic statistical software licenses."),
        ("Pros: ", "Statistically significant. We can confidently say '75 percent of mid-level employees blame poor training'. Forms solid graphs."),
        ("Cons: ", "Surveys are dry. A person might answer '3 out of 5', but we will not know exactly 'why' they gave that score.")
    ]
    for title, text in details2:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(text)

    doc.add_paragraph('\n')
    add_heading(doc, 'OPTION 3: COMPLEX SCOPE')
    details3 = [
        ("Summary: ", "A state-wide, massive study using both massive surveys and direct IT server data. This creates an Academic-grade structural model that proves exactly which factor (e.g., training vs. hardware) causes usage to drop."),
        ("Sampling Method: ", "Multi-stage Cluster Sampling. We divide the entire state of Madhya Pradesh into zones, randomly pick districts from each zone, and survey over 1000+ government employees. We combine this with Focus Group Discussions (10 people in a room) to understand the data deeply."),
        ("Collection Tools: ", "Digital surveys, Focus Group recordings, PLUS secondary IT Server Logs (e.g., extracting login frequencies directly from the NIC/e-Office servers to see real usage, rather than relying on what people tell us)."),
        ("Sorting & Mathematical Algorithms: ", "1. Structural Equation Modeling (SEM) using SmartPLS software. This complex algorithm proves casualty. It maps exactly how much Social Pressure impacts Intention to Use.\n2. K-Means Clustering (Machine Learning): A Python-based sorting algorithm that groups the 1000+ employees into automated 'User Personas' (e.g., 'The Resistant Veteran', 'The Eager Beginner') based on their data patterns."),
        ("Time & Commitment: ", "5 to 6 Months. High cost. Requires high-level permissions to access IT login data. Requires an expert data analyst on the team."),
        ("Pros: ", "Absolute highest standard. Unbeatable evidence. Very impressive to senior leadership and policy makers."),
        ("Cons: ", "Very expensive. Requires massive coordination. If IT server data is denied due to privacy, the model breaks.")
    ]
    for title, text in details3:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(text)

    doc.add_page_break()
    add_heading(doc, 'EXECUTIVE SUMMARY FOR APPROVAL')
    p = doc.add_paragraph()
    p.add_run('Recommendation for the Manager:\n').bold = True
    p.add_run(
        "For an 8-week internship or short-term project, OPTION 1 (Basic Scope) is the most "
        "realistic. It allows you to gather meaningful data quickly without requiring complex "
        "permissions or software. It is also what the current proposal is built around.\n\n"
        "If the department can securely send out an email to thousands of employees and you have "
        "an extra month, OPTION 2 (Moderate Scope) provides better visual graphs and stronger evidence.\n\n"
        "OPTION 3 is typically reserved for year-long government funded research projects or PhD theses."
    )

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Option 1: Basic'
    hdr_cells[2].text = 'Option 2: Moderate'
    hdr_cells[3].text = 'Option 3: Complex'

    rows = [
        ("Primary Method", "Interviews", "Online Surveys", "Mixed + Server Logs"),
        ("Target Audience", "45 to 60", "300 to 400", "1000+ employees"),
        ("Algorithms Used", "Thematic Coding", "SPSS / T-Tests", "SEM & ML Clustering"),
        ("Sampling Type", "Convenience", "Stratified Random", "Multi-stage Cluster"),
        ("Timeline", "8 Weeks", "12-14 Weeks", "5-6 Months"),
        ("Cost & Access", "Very Low", "Medium", "Very High")
    ]
    for row in rows:
        row_cells = table.add_row().cells
        for i in range(4):
            row_cells[i].text = row[i]

    doc.save(os.path.join(BASE, 'Methodology_Scope_Options.docx'))
    print("Saved Methodology_Scope_Options.docx")

if __name__ == '__main__':
    build_final_proposal()
    build_methodology_options()
