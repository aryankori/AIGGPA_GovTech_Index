"""Generate Research Plan DOCX"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC4, 0x9A, 0x2A)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)

# -- Style helpers --
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, italic=False, size=11, color=DARK, spacing_after=6, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(spacing_after)
    if align: p.alignment = align
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1B2A4A')
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ============================================================
# COVER / TITLE
# ============================================================
for _ in range(6): doc.add_paragraph()
add_para('RESEARCH PLAN', bold=True, size=28, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8)
doc.add_paragraph().add_run('━' * 60).font.color.rgb = GOLD
add_para('Assessment of the Use of Digital Tools and Technologies\nby Government Employees for Enhancing\nWorkplace Efficiency and Effectiveness',
         bold=True, size=14, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=20)
doc.add_paragraph().add_run('━' * 60).font.color.rgb = GOLD
for _ in range(4): doc.add_paragraph()

info_lines = [
    ('Prepared By:', 'Aryan Kori'),
    ('Institution:', 'Atal Bihari Vajpayee Institute of Good Governance\nand Policy Analysis (AIGGPA), Bhopal'),
    ('Departments:', 'Revenue, Rural Development, Forest, Health'),
    ('Date:', 'April 2026'),
]
for label, val in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + '  ')
    r1.font.size = Pt(11); r1.bold = True; r1.font.name = 'Calibri'; r1.font.color.rgb = NAVY
    r2 = p.add_run(val)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'; r2.font.color.rgb = DARK

doc.add_page_break()

# ============================================================
# PART 1: KEY OBJECTIVES
# ============================================================
add_heading_styled('PART 1 — KEY OBJECTIVES', 1)

objectives = [
    ('Objective 1: Digital Infrastructure, Awareness & Usage Assessment',
     'What is the current status of digital infrastructure, and to what extent are government employees aware of and using digital tools?',
     ['Inventory of digital devices per employee (desktops, laptops, tablets, smartphones)',
      'Internet bandwidth and uptime metrics at Head Office and District Office levels',
      'Employee awareness levels across a checklist of tools (e-Office, CPGRAMS, departmental portals, email systems)',
      'Frequency of daily/weekly digital tool usage',
      'Observational field data on physical infrastructure conditions']),
    ('Objective 2: Capacity Building & Training Assessment',
     'What capacity building efforts, training programs, and technical support are being provided?',
     ['Number and type of formal IT training programs conducted in the last 2 years',
      'Percentage of employees who attended training, segmented by cadre (Class I\u2013IV)',
      'Duration, content, and format of training programs',
      'Availability of IT helpdesk / on-site technical support',
      'Employee satisfaction scores on training quality and relevance',
      'Mission Karmayogi (iGOT) enrolment and completion records']),
    ('Objective 3: Bottlenecks, Issues & Challenges Identification',
     'What are the major bottlenecks and challenges employees face while using digital tools?',
     ['Types and frequency of reported technical issues (connectivity, hardware, software failures)',
      'Perceived difficulty levels (Effort Expectancy \u2014 Likert scale)',
      'Attitudinal barriers (resistance to change, digital anxiety, generational gaps)',
      'Organizational barriers (unclear mandates, lack of accountability, missing SOPs)',
      'Focus group narratives on systemic pain points',
      'IT complaint/ticket logs (where available)']),
    ('Objective 4: Recommendations for Efficient Service Delivery',
     'What actions can improve the use of digital tools for effective governance?',
     ['Employee-ranked priority areas for improvement',
      'Gap analysis outputs (current state vs. desired state)',
      'Benchmarking data against OECD (2014) and UN DESA (2024) digital government standards',
      'Feasibility and impact assessment of proposed interventions',
      'Alignment check with existing national frameworks (Digital India, Mission Karmayogi)']),
]

for title, question, data_items in objectives:
    add_heading_styled(title, 2)
    p = doc.add_paragraph()
    r = p.add_run('Question this answers: ')
    r.bold = True; r.font.size = Pt(10.5); r.font.name = 'Calibri'; r.font.color.rgb = NAVY
    r2 = p.add_run(question)
    r2.font.size = Pt(10.5); r2.font.name = 'Calibri'; r2.font.color.rgb = DARK
    add_para('Data / Information Needed:', bold=True, size=10.5, color=NAVY, spacing_after=4)
    for item in data_items:
        add_bullet(item)

doc.add_page_break()

# ============================================================
# PART 2: RESEARCH METHODS
# ============================================================
add_heading_styled('PART 2 \u2014 RESEARCH METHODS', 1)

add_heading_styled('2.1 Overall Research Design', 2)
add_para('This study adopts a Convergent Parallel Mixed-Methods Design (Creswell, 2014), wherein quantitative and qualitative data are collected concurrently, analysed independently, and then merged during interpretation to provide a comprehensive understanding.')
add_para('Justification: A mixed-methods approach is particularly appropriate for public administration research because quantitative data reveals the \u201cwhat\u201d (extent of digital tool usage, infrastructure gaps), while qualitative data explains the \u201cwhy\u201d (attitudes, organisational barriers, lived experience of employees) (Creswell, 2014; Heeks, 2006).', italic=True, size=10.5)

add_heading_styled('2.2 Quantitative Methods', 2)
add_heading_styled('2.2.1 Instrument: Structured Schedule', 3)
add_para('A structured schedule will be administered to 320 respondents across 4 departments (Revenue, Rural Development, Forest, Health), stratified by:')
add_bullet('Office type: Head Office + District Office')
add_bullet('Cadre: Class I through Class IV (10 per class per office)')
add_bullet('Age group: 30\u201345 years and 46\u201360 years')
add_para('The schedule will contain:')
add_bullet('Closed-ended items: Device availability checklists, tool awareness checklists, training attendance (Yes/No)')
add_bullet('Likert-scale items (5-point): Internet quality, training satisfaction, perceived difficulty, perceived usefulness')
add_bullet('Frequency items: Digital tool usage (Daily / Weekly / Monthly / Rarely / Never)')

add_heading_styled('2.2.2 Quantitative Analysis Tools', 3)
quant_tools = [
    ['Descriptive Statistics\n(Frequency, Mean, Median, SD)', 'Summarise central tendency and spread', 'Device counts, usage frequency distributions, training attendance rates', 'Field (2018)'],
    ['Cronbach\u2019s Alpha\n(Reliability Testing)', 'Ensure internal consistency of Likert-scale constructs', 'Test reliability of PE, EE, FC sub-scales (threshold: \u03b1 \u2265 0.70)', 'Cronbach (1951);\nTavakol & Dennick (2011)'],
    ['Cross-Tabulation with\nChi-Square Test', 'Examine association between categorical variables', 'Compare tool usage across departments, cadres, age groups', 'Field (2018);\nAgresti (2018)'],
    ['Mann-Whitney U Test', 'Compare ordinal data between two independent groups', 'Compare Likert responses: Head Office vs. District Office', 'Conover (1999);\nField (2018)'],
    ['Kruskal-Wallis H Test', 'Compare ordinal data across three or more groups', 'Compare difficulty levels across four departments; satisfaction across cadres', 'Conover (1999);\nGibbons & Chakraborti (2010)'],
    ['Percentage & Proportion\nAnalysis', 'Quantify awareness and adoption rates', '\u201cX% of Class III employees are unaware of e-Office\u201d', '\u2014'],
]
make_table(['Tool / Technique', 'Purpose', 'Application in This Study', 'Citation'], quant_tools, [1.8, 1.5, 2.2, 1.5])

add_para('')
add_para('Software: SPSS v26+ or equivalent (MS Excel for basic descriptive statistics; SPSS for inferential tests).', bold=True, size=10, color=GRAY)
add_para('Note on Likert Data: Since Likert-scale responses produce ordinal data, non-parametric tests (Mann-Whitney U, Kruskal-Wallis) are preferred over parametric alternatives to avoid violating distributional assumptions (Jamieson, 2004).', italic=True, size=10, color=GRAY)

doc.add_page_break()

add_heading_styled('2.3 Qualitative Methods', 2)
add_heading_styled('2.3.1 Semi-Structured Interviews', 3)
add_para('Interviews will be conducted with a purposively selected sub-sample across all four cadre classes, covering lived experiences with digital tools, perceived barriers and facilitators, training effectiveness narratives, and suggestions for improvement. All interviews will be audio-recorded (with consent), transcribed verbatim, and anonymised.')

add_heading_styled('2.3.2 Focus Group Discussions (FGDs)', 3)
add_para('FGDs will be conducted at the department level to capture collective perspectives on systemic challenges, training adequacy, and policy-level recommendations. Each FGD will comprise 6\u20138 participants from mixed cadres within the same department.')

add_heading_styled('2.3.3 Qualitative Analysis Tools', 3)
qual_tools = [
    ['Reflexive Thematic Analysis\n(6-Phase)', 'Systematically identify, analyse, and report patterns (themes)', 'Primary method for coding interview and FGD transcripts.\n6 phases: (1) Familiarisation, (2) Initial codes, (3) Searching themes, (4) Reviewing, (5) Defining & naming, (6) Report', 'Braun & Clarke\n(2006, 2019)'],
    ['Deductive Coding Frame\n(TAM/UTAUT)', 'Map qualitative findings to established theoretical constructs', 'Code narratives against: Performance Expectancy, Effort Expectancy, Social Influence, Facilitating Conditions', 'Davis (1989);\nVenkatesh et al. (2003)'],
    ['Content Analysis', 'Make replicable inferences from textual data', 'Analyse government circulars, training guidelines, policy documents for stated vs. actual implementation', 'Krippendorff (2018)'],
]
make_table(['Tool / Technique', 'Purpose', 'Application in This Study', 'Citation'], qual_tools, [1.8, 1.5, 2.5, 1.2])

add_para('')
add_para('Coding Approach: A hybrid inductive-deductive approach will be used. Initial codes will emerge inductively from the data, then be mapped to TAM/UTAUT constructs (deductive) to maintain theoretical coherence.', size=10.5)

add_heading_styled('2.4 Observational Method', 2)
add_para('During office visits, a standardised observation checklist will document:')
add_bullet('Physical infrastructure conditions (hardware availability, workspace setup)')
add_bullet('Internet connectivity tests (speed, uptime)')
add_bullet('Actual workflow practices (digital vs. paper-based processes)')
add_bullet('IT support availability and responsiveness')
add_para('This observational data will be triangulated with survey and interview responses to validate self-reported information.')

add_heading_styled('2.5 Gap Analysis & Benchmarking (Objective 4)', 2)
gap_tools = [
    ['Gap Analysis\n(Current vs. Desired State)', 'Identify discrepancies between existing capabilities and targets', 'Compare findings from Obj 1\u20133 against OECD DGPF dimensions and UN EGDI components', 'OECD (2014, 2020)'],
    ['Priority Matrix\n(Impact vs. Feasibility)', 'Rank interventions by potential impact and feasibility', 'Classify recommendations into Quick Wins, Strategic Projects, Fill-ins, Hard Slogs', '\u2014'],
    ['International Benchmarking', 'Contextualise findings within global digital government progress', 'Map outcomes to UN EGDI sub-indices (OSI, TII, HCI) and OECD DGI dimensions', 'United Nations (2024);\nOECD (2014)'],
]
make_table(['Tool / Technique', 'Purpose', 'Application', 'Citation'], gap_tools, [1.8, 1.5, 2.5, 1.2])

doc.add_page_break()

# ============================================================
# PART 3: EVALUATION CRITERIA
# ============================================================
add_heading_styled('PART 3 \u2014 EVALUATION CRITERIA', 1)

add_heading_styled('3.1 Quantitative Evaluation Metrics', 2)
quant_metrics = [
    ['Digital Device Availability Ratio', '\u2265 1 device per employee', 'Count of devices \u00f7 number of employees'],
    ['Internet Connectivity Score', 'Median Likert \u2265 3 (\u201cAdequate\u201d)', '5-point Likert scale'],
    ['Digital Tool Awareness Rate', '\u2265 70% aware of core tools', '% selecting each tool on checklist'],
    ['Daily Digital Tool Usage Rate', '\u2265 50% use tools daily', '% responding \u201cDaily\u201d'],
    ['Training Coverage', '\u2265 60% trained in last 2 years', '% responding \u201cYes\u201d to attendance'],
    ['Training Satisfaction Score', 'Median Likert \u2265 3.5', '5-point Likert scale'],
    ['Perceived Difficulty (Effort Expectancy)', 'Median Likert \u2264 2.5', '5-point Likert (reverse scored)'],
    ['IT Support Availability', '\u2265 70% report accessible support', '% responding \u201cYes\u201d'],
    ['Technical Issue Frequency', '\u2264 20% report daily issues', '% responding \u201cDaily\u201d'],
    ['Instrument Reliability (Cronbach\u2019s \u03b1)', '\u03b1 \u2265 0.70 per construct', 'Cronbach\u2019s Alpha on sub-scales'],
]
make_table(['Metric', 'Benchmark / Threshold', 'Measurement'], quant_metrics, [2.2, 1.8, 3.0])

add_para('')
add_heading_styled('3.2 Qualitative Evaluation Factors', 2)
qual_factors = [
    ['Thematic Saturation', 'No new significant themes emerge in the final 2\u20133 interviews/FGDs'],
    ['Triangulation Convergence', 'Qualitative themes corroborate quantitative findings across \u2265 2 data sources'],
    ['TAM/UTAUT Construct Coverage', 'All four constructs (PE, EE, SI, FC) addressed by at least one theme'],
    ['Actionability of Findings', 'Each bottleneck mapped to a specific, feasible recommendation'],
    ['Participant Representativeness', 'Sub-sample covers all cadres, office types, age groups, and departments'],
]
make_table(['Factor', 'What Constitutes "Success"'], qual_factors, [2.5, 4.5])

add_para('')
add_heading_styled('3.3 Criteria for Viability of Recommendations', 2)
viability = [
    ('Evidence Base', 'Is it directly supported by data from this study?'),
    ('Policy Alignment', 'Does it align with Digital India / Mission Karmayogi objectives?'),
    ('Feasibility', 'Can it be implemented within existing institutional constraints and budget?'),
    ('Scalability', 'Can it be replicated across other departments or states?'),
    ('Impact Potential', 'Will it demonstrably improve service delivery timelines or employee efficiency?'),
]
for i, (criterion, desc) in enumerate(viability, 1):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{i}. {criterion}: ')
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.name = 'Calibri'; r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5); r2.font.name = 'Calibri'; r2.font.color.rgb = DARK

doc.add_page_break()

# ============================================================
# PART 4: EXPECTED OUTCOMES
# ============================================================
add_heading_styled('PART 4 \u2014 EXPECTED OUTCOMES', 1)

add_heading_styled('4.1 Anticipated Findings', 2)

findings = [
    ('Objective 1 (Infrastructure & Awareness)', [
        'Significant disparity in device availability between Head Offices (better equipped) and District Offices',
        'High awareness of basic tools (email) but low awareness of specialised platforms (e-Office, CPGRAMS)',
        'Usage frequency likely correlated with age group (lower in 46\u201360 cohort)',
    ]),
    ('Objective 2 (Capacity Building)', [
        'Formal training coverage likely below 50% of total staff',
        'Training satisfaction likely varies significantly by cadre \u2014 Class III/IV staff may report lower satisfaction',
        'IT helpdesk availability expected to be significantly higher at Head Offices than District Offices',
    ]),
    ('Objective 3 (Bottlenecks)', [
        'Top bottleneck anticipated: poor internet connectivity at District Offices (Facilitating Conditions)',
        'Significant attitudinal resistance among senior employees (46\u201360 age group) \u2014 Effort Expectancy',
        'Lack of departmental mandate / accountability for digital adoption \u2014 Social Influence',
        'Software usability issues (complex interfaces) \u2014 Performance Expectancy',
    ]),
    ('Objective 4 (Recommendations)', [
        'Infrastructure investment priorities (bandwidth, devices) for District Offices',
        'Cadre-specific, role-based training programmes (not one-size-fits-all)',
        'Establishment of permanent IT helpdesk at District level',
        'Policy mandate for digital-first workflows with transition support',
    ]),
]

for title, items in findings:
    add_para(title, bold=True, size=11, color=NAVY, spacing_after=4)
    for item in items:
        add_bullet(item)

add_heading_styled('4.2 Next Steps Following Research', 2)
next_steps = [
    ['Immediate (Week 4)', 'Present findings and evidence-based recommendations to AIGGPA senior management', 'End of internship'],
    ['Short-term (1\u20133 months)', 'Department-level dissemination; initiate pilot training programmes', 'Post-internship'],
    ['Medium-term (3\u20136 months)', 'Policy brief prepared for state administration', 'Post-internship'],
    ['Long-term (6\u201312 months)', 'Follow-up study to measure impact of implemented recommendations', 'Future research'],
]
make_table(['Phase', 'Action', 'Timeline'], next_steps, [1.8, 3.5, 1.7])

add_para('')
add_heading_styled('4.3 Deliverables', 2)
deliverables = [
    'Final Research Report \u2014 comprehensive mixed-methods analysis with all findings, visualisations, and recommendations',
    'Executive Summary \u2014 2-page brief for senior management',
    'Policy Brief \u2014 actionable recommendations aligned with Digital India / Mission Karmayogi',
    'Raw Data Archive \u2014 anonymised datasets for institutional records',
    'Presentation Deck \u2014 visual presentation for stakeholder briefing',
]
for i, d in enumerate(deliverables, 1):
    add_bullet(f'{i}. {d}')

doc.add_page_break()

# ============================================================
# REFERENCES
# ============================================================
add_heading_styled('REFERENCES', 1)

refs = [
    'Agresti, A. (2018). Statistical Methods for the Social Sciences (5th ed.). Pearson.',
    'Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77\u2013101.',
    'Braun, V., & Clarke, V. (2019). Reflecting on reflexive thematic analysis. Qualitative Research in Sport, Exercise and Health, 11(4), 589\u2013597.',
    'Conover, W. J. (1999). Practical Nonparametric Statistics (3rd ed.). Wiley.',
    'Creswell, J. W. (2014). Research Design: Qualitative, Quantitative, and Mixed Methods Approaches (4th ed.). SAGE Publications.',
    'Cronbach, L. J. (1951). Coefficient alpha and the internal structure of tests. Psychometrika, 16(3), 297\u2013334.',
    'Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319\u2013340.',
    'Field, A. (2018). Discovering Statistics Using IBM SPSS Statistics (5th ed.). SAGE Publications.',
    'Gibbons, J. D., & Chakraborti, S. (2010). Nonparametric Statistical Inference (5th ed.). CRC Press.',
    'Government of India, Department of Personnel & Training. (2020). National Programme for Civil Services Capacity Building (Mission Karmayogi).',
    'Government of India, Ministry of Electronics & Information Technology. (2015). Digital India.',
    'Heeks, R. (2003). Most eGovernment-for-Development Projects Fail: How Can Risks Be Reduced? University of Manchester.',
    'Heeks, R. (2006). Implementing and Managing eGovernment. SAGE Publications.',
    'Jamieson, S. (2004). Likert scales: How to (ab)use them. Medical Education, 38(12), 1217\u20131218.',
    'Krippendorff, K. (2018). Content Analysis: An Introduction to Its Methodology (4th ed.). SAGE Publications.',
    'OECD. (2014). Recommendation of the Council on Digital Government Strategies. OECD Publishing.',
    'OECD. (2020). Digital Government Index: 2019 Results. OECD Public Governance Policy Papers, No. 03.',
    'Roni, S. M., & Djajadikerta, H. G. (2021). Data Analysis with SPSS for Survey-based Research. Springer.',
    'Tavakol, M., & Dennick, R. (2011). Making sense of Cronbach\u2019s alpha. International Journal of Medical Education, 2, 53\u201355.',
    'United Nations. (2024). E-Government Survey 2024: The Future of Digital Government. UN DESA.',
    'Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425\u2013478.',
    'Venkatesh, V., Thong, J. Y. L., & Xu, X. (2012). Consumer acceptance and use of information technology. MIS Quarterly, 36(1), 157\u2013178.',
    'World Bank. (2016). World Development Report 2016: Digital Dividends. The World Bank.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

# ============================================================
# SAVE
# ============================================================
out = 'AIGGPA_Research_Plan.docx'
doc.save(out)
print(f'Done: {out}')
